from docx import Document as Doc
from docx.document import Document
from docx.section import Section
from docx.text.paragraph import Paragraph
from docx.text.parfmt import ParagraphFormat
from docx.text.run import Font
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.text.run import Run

from abc import ABC, abstractmethod
from docx.shared import Mm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from typing import Any


class RunStyle(ABC):

    _run: Run
    _paragraph: Paragraph

    def __init__(self, paragraph: Paragraph):
        self._paragraph = paragraph
        self._run = self._paragraph.add_run()
        self.__set_default_run_font()
    
    def __del__(self):
        if self._paragraph != None and self._run != None:
            self._paragraph._p.remove(self._run._r)

    def __set_default_run_font(self) -> None:
        #WD_COLOR_INDEX for highlight_color, not rgb
        self._run.font.name = "Times New Roman"
        self._run.font.size = Pt(14)
        self._run.font.highlight_color = None
        self._run.font.color.rgb = RGBColor(0, 0, 0)
        self._run.font.bold = False
        self._run.font.italic = False
        self._run.font.underline = False

    def _set_run_font(self, dst_run_font: Font, 
                      src_run_font: Font) -> None:
        dst_run_font.name = src_run_font.highlight_color
        dst_run_font.size = src_run_font.size
        dst_run_font.highlight_color = src_run_font.highlight_color
        dst_run_font.color.rgb = src_run_font.color.rgb
        dst_run_font.bold = src_run_font.bold
        dst_run_font.italic = src_run_font.italic
        dst_run_font.underline = src_run_font.underline

    def convert_run_font(self, run: Run) -> None:
        self._set_run_font(run.font, self._run.font)


class CommonParagraphRun(RunStyle):

    def __init__(self, paragraph: Paragraph):
        super().__init__(paragraph)


class ParagraphStyle(ABC):

    _paragraph: Paragraph = None
    _paragraph_style_format: ParagraphFormat = None
    _paragraph_style_font: Font = None
    _style: Any = None

    def __init__(self, document: Document, style_name: str,
                 next_par_style: Any, indent: Mm,
                 alignment: WD_ALIGN_PARAGRAPH):
        if style_name not in document.styles:
            document.styles.add_style(style_name, 
                                      WD_STYLE_TYPE.PARAGRAPH)
            self._paragraph = document.add_paragraph("Some text",
                                                     style_name)
            self._style = document.styles[style_name]
            self._paragraph_style_font = self._style.font
            self._paragraph_style_format = \
                self._style.paragraph_format
            self.__set_default_paragraph_style(next_par_style, 
                                               indent,
                                               alignment)
        else:
            self._style = document.styles[style_name]
            self._paragraph_style_font = self._style.font
            self._paragraph_style_format = \
            self._style.paragraph_format
            self._paragraph = document.add_paragraph("Some text",
                                                     style_name)

    def __del__(self):
        pass
        #self.__delete_paragraph()
    
    def __delete_paragraph(self):
        if self._paragraph != None:
            p = self._paragraph._element
            p.getparent().remove(p)
            self._paragraph._element = None
            self._paragraph._p = None

    def __set_default_paragraph_style_font(self) -> None:       
        self._paragraph_style_font.name = "Times New Roman"
        self._paragraph_style_font.size = Pt(14)
        self._paragraph_style_font.highlight_color = None
        self._paragraph_style_font.color.rgb = RGBColor(0, 0, 0)
        self._paragraph_style_font.bold = False
        self._paragraph_style_font.italic = False
        self._paragraph_style_font.underline = False

    def __set_default_paragraph_style_format(
            self, indent: Mm, alignment: WD_ALIGN_PARAGRAPH) -> None:
        self._paragraph_style_format.first_line_indent = indent
        self._paragraph_style_format.alignment = alignment
        self._paragraph_style_format.space_before = Mm(0)
        self._paragraph_style_format.space_after = Mm(0)
        self._paragraph_style_format.line_spacing_rule = \
            WD_LINE_SPACING.MULTIPLE
        self._paragraph_style_format.line_spacing = 1.5

    def __set_default_paragraph_style(self, next_par_style: Any,
                                      indent: Mm,
                                      alignment: WD_ALIGN_PARAGRAPH
                                      ) -> None:
        if next_par_style == None:
            self._style.next_paragraph_style = self._style
        else:
            self._style.next_paragraph_style = next_par_style
        self.__set_default_paragraph_style_font()
        self.__set_default_paragraph_style_format(indent, alignment)

    def _set_style_font(self, dst_style_font: Font, 
                        src_style_font: Font) -> None:
        dst_style_font.name = src_style_font.name
        dst_style_font.size = src_style_font.size
        dst_style_font.highlight_color = \
            src_style_font.highlight_color
        dst_style_font.color.rgb = src_style_font.color.rgb
        dst_style_font.bold = src_style_font.bold
        dst_style_font.italic = src_style_font.bold
        dst_style_font.underline = src_style_font.bold    

    def _set_paragraph_style_format(self, \
                                    dst_par_format: ParagraphFormat,
                                    src_par_format: ParagraphFormat
                                    ) -> None:
        dst_par_format \
            .first_line_indent = src_par_format.first_line_indent
        dst_par_format.alignment = src_par_format.alignment
        dst_par_format.space_before = src_par_format.space_before
        dst_par_format.space_after = src_par_format.space_after
        dst_par_format \
            .line_spacing_rule = src_par_format.line_spacing_rule
        dst_par_format.line_spacing = src_par_format.line_spacing

    def delete_tmp_paragraph(self):
        self.__delete_paragraph()

    def convert_paragraph_instance_style(self, paragraph: Paragraph):
        paragraph.style = self._paragraph.style.name
        paragraph.paragraph_format.first_line_indent = \
            self._paragraph.paragraph_format.first_line_indent

class CommonTextParagraphStyle(ParagraphStyle):

    def __init__(self, document: Document):
        super().__init__(document, "common_text_paragraph_style", 
                         None, Mm(12.5), WD_ALIGN_PARAGRAPH.JUSTIFY)
        
    def __del__(self):
        super().__del__()
        
class CommonPictureCaptionStyle(ParagraphStyle):

    def __init__(self, document: Document):
        super().__init__(document, "common_picture_caption_style",
                         None, Mm(0), WD_ALIGN_PARAGRAPH.CENTER)
    
    def __del__(self):
        super().__del__()
        
class CommonTableCaptionStyle(ParagraphStyle):

    def __init__(self, document: Document):
        super().__init__(document, "common_table_caption_style",
                         None, Mm(0), WD_ALIGN_PARAGRAPH.RIGHT)
        
    def __del__(self):
        super().__del__()
        
class ParagraphStyleCreator(ABC):

    _document: Document

    def __init__(self, document: Document):
        self._document = document

    @abstractmethod
    def get_paragraph_style() -> ParagraphStyle:
        pass

class TextParagraphStyleCreator(ParagraphStyleCreator):

    def __init__(self, document: Document):
        super().__init__(document)

    def get_paragraph_style(self, text_paragraph_style_name: str
                            ) -> ParagraphStyle:
        if (text_paragraph_style_name 
            == "common_text_paragraph_style"):
            return CommonTextParagraphStyle(self._document)
        return None

class PictureCaptionCreator(ParagraphStyleCreator):

    def __init__(self, document: Document):
        super().__init__(document)

    def get_paragraph_style(self, picture_captinon_style_name: str
                            ) -> ParagraphStyle:
        if (picture_captinon_style_name 
            == "common_picture_caption_style"):
            return CommonPictureCaptionStyle(self._document)
        return None

class TableCaptionCreator(ParagraphStyleCreator):

    def __init__(self, document: Document):
        super().__init__(document)

    def get_paragraph_style(self, table_captinon_style_name: str
                            ) -> ParagraphStyle:
        if (table_captinon_style_name 
            == "common_table_caption_style"):
            return CommonTableCaptionStyle(self._document)
        return None